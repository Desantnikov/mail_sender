import re

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TextPartToEnter:
    def __init__(self, paragraph):
        self.paragraph = paragraph

        self.text = paragraph.text
        self.alignment = WD_ALIGN_PARAGRAPH._member_to_xml.get(paragraph.alignment, '').upper()

    def __str__(self):
        return self.text

    def fill_hyperlink(self, relation_id, relation_target):
        """ w/a to handle hyperlinks
            see MailTemplate._fill_hyperlinks() docstring for details"""

        # find position in xml where relation is declared
        relation_index = self.paragraph._element.xml.find(relation_id)

        # -1 if was not found
        if relation_index < 0:
            return

        # take 100 previous symbols of xml file
        previous_xml_part = self.paragraph._element.xml[relation_index-100:relation_index]

        # look for normal (not xml-related) text
        text_before_link = re.search(r'>([А-яA-z0-9]+:\s*)<', previous_xml_part).group(1)

        # index of the end of text before link (assume that link's text should start here)
        text_before_link_end_index = self.text.index(text_before_link) + len(text_before_link)

        # insert link's text at the place where it should be
        self.text = f'{self.text[:text_before_link_end_index]}{relation_target}{self.text[text_before_link_end_index:]}'


class MailTemplate:
    def __init__(self, doc: Document, attachment_files_sequence):
        self.doc = doc
        self.paragraphs = doc.paragraphs

        self.topic = None
        self.hyperlinks = None
        self.body = []
        self.attachment_files_sequence = attachment_files_sequence

        self._set_theme()
        self._set_body()

        self._find_hyperlinks()
        self._fill_hyperlinks()

    def __str__(self):
        return f'Тема:\r\n{str(self.topic)}\r\nТекст:{self._get_body_as_plain_text()}\r\n'

    def _set_theme(self):
        self.topic = self.paragraphs[0].text.replace('Тема: ', '').strip()

    def _set_body(self):
        # 0-th line for mail topic
        for index, paragraph in enumerate(self.paragraphs[1:]):
            if paragraph.text:
                starting_paragraph_index = index + 1
                break

        self.body = [TextPartToEnter(paragraph=paragraph) for paragraph in self.paragraphs[starting_paragraph_index:]]



    def _get_body_as_plain_text(self):
        return "\r\n".join([part.text for part in self.body])

    def _find_hyperlinks(self):
        """ part of ugly w/a to handle hyperlinks """
        self.hyperlinks = {relation_id: relation for relation_id, relation in self.doc.part.rels.items() if
                           relation.reltype == RELATIONSHIP_TYPE.HYPERLINK}

    def _fill_hyperlinks(self):
        """ ugly workaround to insert hyperlink's text at the place where it should be
            since python-docx can't handle hyperlinks reading

            takes all hyperlinks and looks for their identifiers in each paragraph's raw xml
            if found tries to insert link's text at the place where it should be """

        for relation_id, relation in self.hyperlinks.items():
            for text_part in self.body:
                text_part.fill_hyperlink(relation_id, relation._target)


class Mail:
    def __init__(self, recipient: str, mail_template: MailTemplate):
        self.recipient = recipient

        self.topic = mail_template.topic
        self.body = mail_template.body
        self.attachment_files_sequence = mail_template.attachment_files_sequence

    def __str__(self):
        return "\r\n".join([part.text for part in self.body])